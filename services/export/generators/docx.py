"""Word generator — uses python-docx. Template optional."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


_EY_CHARCOAL = RGBColor(0x2E, 0x2E, 0x38)
_EY_YELLOW = RGBColor(0xFF, 0xE6, 0x00)


def _style_default_fonts(doc: Document) -> None:
    """Apply EY-style default body font when no template is provided."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = _EY_CHARCOAL


def build_docx(
    output_path: Path,
    plan: dict[str, Any],
    template_path: Path | None = None,
) -> Path:
    """Build a .docx from the LLM plan.

    Plan shape:
      {
        "title": str,
        "subtitle": str,
        "sections": [
          {"heading": str, "paragraphs": [str], "bullets": [str]?}
        ]
      }
    """
    if template_path and template_path.exists():
        doc = Document(str(template_path))
        # Clear existing body content but keep styles
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        doc = Document()
        _style_default_fonts(doc)

    title = plan.get("title") or "Document"
    subtitle = plan.get("subtitle") or ""

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = _EY_CHARCOAL

    if subtitle:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(subtitle)
        sub_run.font.size = Pt(13)
        sub_run.font.italic = True
        sub_run.font.color.rgb = _EY_CHARCOAL

    doc.add_paragraph()  # spacer

    for section in plan.get("sections") or []:
        heading = section.get("heading") or ""
        if heading:
            h = doc.add_paragraph()
            h_run = h.add_run(heading)
            h_run.font.size = Pt(14)
            h_run.font.bold = True
            h_run.font.color.rgb = _EY_CHARCOAL

        for para in section.get("paragraphs") or []:
            doc.add_paragraph(para)

        for bullet in section.get("bullets") or []:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.save(output_path)
    return output_path
